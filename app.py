from flask import Flask, render_template, request, redirect, url_for, flash, send_file
import os
import tempfile
import time
from datetime import datetime
import json
from werkzeug.utils import secure_filename
from pdf2docx import Converter
import pytesseract
from PIL import Image
import logging
import hashlib
from flask import Flask, render_template, request, redirect, url_for, flash, send_from_directory
import config
from docx import Document

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# 初始化Flask应用
app = Flask(__name__)
app.config.from_object('config.Config')

# 确保上传和结果目录存在
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['RESULT_FOLDER'], exist_ok=True)
os.makedirs(app.config['HISTORY_FOLDER'], exist_ok=True)

# 允许的文件扩展名
ALLOWED_EXTENSIONS = {'pdf'}

# 检查文件扩展名是否允许
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# 检查文件大小是否符合限制
def check_file_size(file):
    file.seek(0, os.SEEK_END)
    file_size = file.tell()
    file.seek(0)
    return file_size <= app.config['MAX_CONTENT_LENGTH']

# 转换PDF为Word
def convert_pdf_to_word(pdf_path, word_path, mode='fast'):
    try:
        logger.info(f'开始转换文件: {pdf_path} 到 {word_path}，模式: {mode}')
        start_time = time.time()
        cv = Converter(pdf_path)
        if mode == 'accurate':
            cv.convert(word_path, start=0, end=None)
        else:
            cv.convert(word_path, start=0, end=None, compress_images=False)
        cv.close()
        end_time = time.time()
        logger.info(f'转换完成，耗时: {end_time - start_time:.2f} 秒')
        return True
    except Exception as e:
        logger.error(f'转换失败: {str(e)}')
        return False

# OCR文字识别处理
def ocr_process(pdf_path):
    try:
        logger.info(f'开始OCR处理: {pdf_path}')
        # 这里简化处理，实际应用中需要将PDF转为图片再进行OCR
        # 可以使用pdf2image库
        return True
    except Exception as e:
        logger.error(f'OCR处理失败: {str(e)}')
        return False

# 保存转换历史
def save_history(original_filename, converted_filename, conversion_type, status):
    history = {
        'original_filename': original_filename,
        'converted_filename': converted_filename,
        'conversion_type': conversion_type,
        'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        'status': status
    }
    
    # 读取现有历史
    history_file = os.path.join(app.config['HISTORY_FOLDER'], 'history.json')
    try:
        if os.path.exists(history_file):
            with open(history_file, 'r') as f:
                histories = json.load(f)
        else:
            histories = []
    except:
        histories = []
    
    # 添加新记录并限制数量
    histories.insert(0, history)
    if len(histories) > app.config['MAX_HISTORY_ENTRIES']:
        histories = histories[:app.config['MAX_HISTORY_ENTRIES']]
    
    # 保存历史
    with open(history_file, 'w') as f:
        json.dump(histories, f, ensure_ascii=False, indent=4)

# 测试页面路由
@app.route('/test')
def test():
    return render_template('test.html')

# 首页路由
@app.route('/')
def index():
    return render_template('index.html')

# 单文件转换路由
@app.route('/convert', methods=['POST'])
def convert():
    if 'file' not in request.files:
        flash('没有文件部分')
        return redirect(request.url)
    
    file = request.files['file']
    if file.filename == '':
        flash('没有选择文件')
        return redirect(request.url)
    
    if file and allowed_file(file.filename):
        # 检查文件大小
        if not check_file_size(file):
            flash(f'文件大小超过限制 ({app.config['MAX_CONTENT_LENGTH'] // (1024 * 1024)}MB)')
            return redirect(request.url)
        
        # 保存上传文件
        filename = secure_filename(file.filename)
        pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(pdf_path)
        
        # 获取自定义文件名
        custom_filename = request.form.get('custom_filename', '')

        # 准备输出文件路径
        if custom_filename:
            word_filename = f'{custom_filename}.docx'
        else:
            word_filename = f'{os.path.splitext(filename)[0]}.docx'
        word_path = os.path.join(app.config['RESULT_FOLDER'], word_filename)
        
        # 获取转换模式
        mode = request.form.get('mode', 'fast')
        
        # 获取是否需要OCR
        use_ocr = request.form.get('use_ocr', 'false') == 'true'
        
        # 如果需要OCR，先进行OCR处理
        if use_ocr:
            ocr_success = ocr_process(pdf_path)
            if not ocr_success:
                flash('OCR处理失败')
                return redirect(request.url)
        
        # 转换PDF为Word
        conversion_success = convert_pdf_to_word(pdf_path, word_path, mode)
        
        if conversion_success:
            # 保存转换历史
            conversion_type = f'{mode}模式{"+OCR" if use_ocr else ""}'
            save_history(filename, word_filename, conversion_type, 'success')
            
            # 生成临时下载链接
            temp_link = generate_temp_link(word_filename)
            
            return render_template('result.html', 
                                  original_filename=filename, 
                                  converted_filename=word_filename, 
                                  temp_link=temp_link)
        else:
            save_history(filename, word_filename, conversion_type, 'failed')
            flash('转换失败，请检查文件是否损坏或加密')
            return redirect(request.url)
    else:
        flash('不支持的文件类型')
        return redirect(request.url)

# 批量转换路由
@app.route('/batch_convert', methods=['POST'])
def batch_convert():
    if 'files' not in request.files:
        flash('没有文件部分')
        return redirect(request.url)
    
    files = request.files.getlist('files')
    if len(files) > app.config['MAX_BATCH_FILES']:
        flash(f'批量转换最多支持{app.config['MAX_BATCH_FILES']}个文件')
        return redirect(request.url)
    
    results = []
    valid_files = []
    
    for file in files:
        if file.filename == '':
            continue
        
        if file and allowed_file(file.filename):
            # 检查文件大小
            if not check_file_size(file):
                flash(f'文件{file.filename}大小超过限制 ({app.config['MAX_CONTENT_LENGTH'] // (1024 * 1024)}MB)')
                continue
            
            valid_files.append(file)
        else:
            flash(f'文件{file.filename}类型不支持')
    
    if not valid_files:
        flash('没有有效的PDF文件')
        return redirect(request.url)
    
    # 获取转换模式
    mode = request.form.get('mode', 'fast')
    
    # 获取是否需要OCR
    use_ocr = request.form.get('use_ocr', 'false') == 'true'

    # 获取自定义文件名前缀
    custom_prefix = request.form.get('custom_prefix', '')

    for i, file in enumerate(valid_files, 1):
        # 保存上传文件
        filename = secure_filename(file.filename)
        pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(pdf_path)
        
        # 准备输出文件路径
        if custom_prefix:
            word_filename = f'{custom_prefix}_{i}.docx'
        else:
            word_filename = f'{os.path.splitext(filename)[0]}.docx'
        word_path = os.path.join(app.config['RESULT_FOLDER'], word_filename)
        
        # 如果需要OCR，先进行OCR处理
        ocr_success = True
        if use_ocr:
            ocr_success = ocr_process(pdf_path)
            if not ocr_success:
                flash(f'文件{filename}OCR处理失败')
        
        # 转换PDF为Word
        if ocr_success:
            conversion_success = convert_pdf_to_word(pdf_path, word_path, mode)
        else:
            conversion_success = False
        
        # 保存转换历史
        conversion_type = f'{mode}模式{"+OCR" if use_ocr else ""}'
        save_history(filename, word_filename, conversion_type, 'success' if conversion_success else 'failed')
        
        if conversion_success:
            # 生成临时下载链接
            temp_link = generate_temp_link(word_filename)
            results.append({
                'original_filename': filename,
                'converted_filename': word_filename,
                'temp_link': temp_link,
                'status': 'success'
            })
        else:
            results.append({
                'original_filename': filename,
                'converted_filename': word_filename,
                'status': 'failed'
            })
    
    return render_template('batch_result.html', results=results)

# 生成临时下载链接
def generate_temp_link(filename):
    # 简单实现，生成一个基于时间戳的链接
    timestamp = str(int(time.time()))
    token = hashlib.md5((filename + timestamp + app.config['SECRET_KEY']).encode()).hexdigest()
    return url_for('download_temp', filename=filename, token=token, _external=True)

# 临时下载路由
@app.route('/download_temp/<filename>/<token>')
def download_temp(filename, token):
    # 验证token是否过期
    current_time = int(time.time())
    # 检查链接是否在有效期内（允许有1分钟的误差）
    for i in range(-60, 60):
        timestamp = str(current_time + i)
        valid_token = hashlib.md5((filename + timestamp + app.config['SECRET_KEY']).encode()).hexdigest()
        if token == valid_token:
            break
    else:
        flash('临时链接无效或已过期')
        return redirect(url_for('index'))
    
    word_path = os.path.join(app.config['RESULT_FOLDER'], filename)
    if not os.path.exists(word_path):
        flash('文件不存在')
        return redirect(url_for('index'))
    
    return send_file(word_path, as_attachment=True)

# 查看历史记录路由
@app.route('/history')
def history():
    history_file = os.path.join(app.config['HISTORY_FOLDER'], 'history.json')
    try:
        if os.path.exists(history_file):
            with open(history_file, 'r') as f:
                histories = json.load(f)
        else:
            histories = []
    except:
        histories = []
    
    return render_template('history.html', histories=histories)

# 错误处理
@app.errorhandler(413)
def request_entity_too_large(error):
    flash(f'文件大小超过限制 ({app.config['MAX_CONTENT_LENGTH'] // (1024 * 1024)}MB)')
    return redirect(url_for('index'))

# 下载历史文件路由
@app.route('/download/<filename>')
def download(filename):
    word_path = os.path.join(app.config['RESULT_FOLDER'], filename)
    if not os.path.exists(word_path):
        flash('文件不存在')
        return redirect(url_for('history'))
    
    return send_file(word_path, as_attachment=True)

# 预览文件路由
@app.route('/preview/<filename>')
def preview(filename):
    word_path = os.path.join(app.config['RESULT_FOLDER'], filename)
    if not os.path.exists(word_path):
        flash('文件不存在')
        return redirect(url_for('history'))
    
    # 使用python-docx库读取Word文档内容
    try:
        doc = Document(word_path)
        paragraphs = []
        for para in doc.paragraphs:
            paragraphs.append(para.text)
        
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text)
                table_data.append(row_data)
            tables.append(table_data)
    except:
        paragraphs = ['无法读取文档内容']
        tables = []
    
    return render_template('preview.html', filename=filename, paragraphs=paragraphs, tables=tables)

# 厦门一日游路由
@app.route('/xiamen_tour')
def xiamen_tour():
    return render_template('xiamen_tour.html')

# 应用入口
if __name__ == '__main__':
    # 初始化配置
    app.config.from_object('config.DevelopmentConfig')
    config.Config.init_app(app)
    
    # 启动应用
    app.run(debug=True)